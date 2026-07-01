from docx import Document as DocxDocument

from app.ingestion.models.document import (
    Document
)

from app.ingestion.parsers.base_parser import (
    BaseParser
)


class DOCXParser(BaseParser):

    def parse(
        self,
        file_path: str
    ) -> Document:

        doc = DocxDocument(file_path)

        text = "\n".join(
            [
                para.text
                for para in doc.paragraphs
            ]
        )

        return Document(
            content=text,
            metadata={},
            source=file_path,
            document_type="docx"
        )